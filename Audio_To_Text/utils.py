import io
from docx import Document

# Function to convert text to a TXT file
def convert_to_txt(response):
    """Convert the text response to a TXT format."""
    try:
        return str(response).encode("utf-8", errors="ignore")
    except Exception as e:
        raise ValueError(f"Error converting to TXT: {e}")

# Function to convert text to a DOCX file
def convert_to_docx(response):
    """Convert the text response to a DOCX format."""
    try:
        doc = Document()
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(response)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    except Exception as e:
        raise ValueError(f"Error converting to DOCX: {e}")
